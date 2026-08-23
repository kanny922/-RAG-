# -*- coding: utf-8 -*-
"""
document_loader.py —— 文档统一解析模块
======================================

本模块是 RAG 流水线的第一站，负责把不同格式的原始文档（PDF / Word / TXT）
统一解析成 :class:`Document` 对象，供后续的文本分块模块使用。

支持能力：
1. ``.pdf``  ：使用 PyMuPDF（import 名为 fitz）逐页提取文本；
               若某页提取不到文本（通常是扫描件 / 图片型 PDF），
               自动进入 OCR 分支（pdf2image 转图片 + rapidocr_onnxruntime 识别）。
2. ``.docx`` ：使用 python-docx 提取全部段落文本。
3. ``.txt``  ：直接按 utf-8 / gbk 编码读取。

设计说明（教学要点）：
- 所有解析结果都封装成统一的 Document，后续模块无需关心文件格式差异；
- OCR 属于"重依赖"，本模块采用**延迟导入 + 优雅降级**：
  没装 OCR 依赖时不会崩溃，而是打印中文警告并跳过该页。
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List


@dataclass
class Document:
    """统一文档对象：所有解析结果都用它承载。

    Attributes:
        page_content: 文档的正文文本。
        metadata:     描述信息字典，至少包含以下键：
                      - source   : 来源文件名（用于答案溯源）
                      - page     : 页码（PDF 有页码；docx / txt 为 None）
                      - doc_type : 文档类型（pdf / docx / txt）
    """

    page_content: str
    metadata: Dict = field(default_factory=dict)

    def __post_init__(self):
        """补全 metadata 的默认字段，保证下游模块取值时不会 KeyError。"""
        self.metadata.setdefault("source", "未知来源")
        self.metadata.setdefault("page", None)
        self.metadata.setdefault("doc_type", "unknown")


class DocumentLoader:
    """文档统一加载器：根据文件扩展名自动选择解析方式。"""

    #: 当前支持的文件扩展名（app.py 会用它扫描知识库目录）
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def __init__(self):
        # OCR 引擎采用懒加载：只有真正遇到扫描件时才初始化（初始化较慢）
        self._ocr_engine = None      # 缓存 RapidOCR 实例，避免每页重复创建
        self._ocr_checked = False    # 标记是否已尝试过加载 OCR 依赖

    # ------------------------------------------------------------------ #
    # 对外接口
    # ------------------------------------------------------------------ #
    def load(self, path) -> List[Document]:
        """加载单个文件，自动按扩展名路由到对应的解析函数。

        Args:
            path: 文件路径（str 或 Path）。

        Returns:
            Document 列表：PDF 每页一个 Document；docx / txt 整体一个 Document。

        Raises:
            FileNotFoundError: 文件不存在。
            ValueError:        不支持的文件格式。
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在：{path}")

        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._load_pdf(path)
        if ext == ".docx":
            return self._load_docx(path)
        if ext == ".txt":
            return self._load_txt(path)
        raise ValueError(
            f"暂不支持的文件格式：{ext}，目前支持：{sorted(self.SUPPORTED_EXTENSIONS)}"
        )

    def load_directory(self, dir_path) -> List[Document]:
        """批量加载目录下的全部支持格式文档（含子目录）。

        单个文件解析失败不会中断整体流程，只会打印警告并跳过——
        知识库里混入一个损坏文件也不应导致整个系统无法启动。

        Args:
            dir_path: 文档目录路径。

        Returns:
            所有文档合并后的 Document 列表。
        """
        dir_path = Path(dir_path)
        if not dir_path.exists():
            raise FileNotFoundError(f"目录不存在：{dir_path}")

        # 递归找出全部支持的文件，按文件名排序保证入库顺序稳定
        files = sorted(
            p for p in dir_path.rglob("*")
            if p.is_file() and p.suffix.lower() in self.SUPPORTED_EXTENSIONS
        )
        if not files:
            print(f"[提示] 目录 {dir_path} 中没有找到可解析的文档（支持 pdf / docx / txt）。")
            return []

        all_documents: List[Document] = []
        for file_path in files:
            try:
                docs = self.load(file_path)
                all_documents.extend(docs)
                print(f"[信息] 已加载《{file_path.name}》：{len(docs)} 个文档片段。")
            except Exception as exc:  # 单个文件失败不影响其他文件
                print(f"[警告] 加载《{file_path.name}》失败，已跳过。原因：{exc}")
        return all_documents

    # ------------------------------------------------------------------ #
    # PDF 解析（含 OCR 降级分支）
    # ------------------------------------------------------------------ #
    def _load_pdf(self, path: Path) -> List[Document]:
        """用 PyMuPDF 逐页提取文本；无文本层的页走 OCR 识别。"""
        try:
            import fitz  # PyMuPDF 的 import 名是 fitz
        except ImportError as exc:
            raise ImportError("缺少 PyMuPDF 依赖，请先执行：pip install pymupdf") from exc

        documents: List[Document] = []
        pdf = fitz.open(str(path))
        try:
            for page_index in range(len(pdf)):
                page = pdf[page_index]
                text = page.get_text("text").strip()

                if not text:
                    # 该页没有文本层 → 大概率是扫描件，尝试 OCR
                    print(f"[提示] 《{path.name}》第 {page_index + 1} 页无文本层，尝试 OCR 识别……")
                    text = self._ocr_pdf_page(path, page_index)

                if text:
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            "source": path.name,
                            "page": page_index + 1,  # 页码从 1 开始，符合阅读习惯
                            "doc_type": "pdf",
                        },
                    ))
                else:
                    print(f"[警告] 《{path.name}》第 {page_index + 1} 页未能提取到任何文本，已跳过。")
        finally:
            pdf.close()  # 无论是否出错都要释放文件句柄
        return documents

    def _get_ocr_engine(self):
        """懒加载 OCR 引擎；依赖缺失时打印中文警告并返回 None（优雅降级）。"""
        if self._ocr_checked:
            return self._ocr_engine
        self._ocr_checked = True
        try:
            from rapidocr_onnxruntime import RapidOCR
            self._ocr_engine = RapidOCR()
            print("[信息] OCR 引擎初始化完成。")
        except ImportError:
            print("[警告] 未安装 rapidocr_onnxruntime，OCR 功能不可用。")
            print("       如需识别扫描版 PDF，请执行：pip install rapidocr_onnxruntime pdf2image")
            self._ocr_engine = None
        except Exception as exc:  # 例如模型文件下载失败等
            print(f"[警告] OCR 引擎初始化失败：{exc}")
            self._ocr_engine = None
        return self._ocr_engine

    def _ocr_pdf_page(self, path: Path, page_index: int) -> str:
        """把 PDF 指定页渲染成图片并做 OCR，返回识别出的文本。

        注意：pdf2image 依赖系统级软件 poppler（Windows 需单独安装并配置 PATH）。
        任何一步失败都只打印警告并返回空字符串，绝不让 OCR 问题拖垮主流程。
        """
        try:
            import numpy as np
            from pdf2image import convert_from_path
        except ImportError:
            print("[警告] 缺少 pdf2image / numpy 依赖，无法执行 OCR，已跳过该页。")
            return ""

        engine = self._get_ocr_engine()
        if engine is None:
            return ""

        try:
            # first_page / last_page 都是 1 起始的页码，只渲染当前这一页
            images = convert_from_path(
                str(path), dpi=200,
                first_page=page_index + 1, last_page=page_index + 1,
            )
            # RapidOCR 返回 (结果列表, 耗时)，每条结果为 [文本框坐标, 文本, 置信度]
            result, _ = engine(np.array(images[0]))
            if not result:
                return ""
            return "\n".join(line[1] for line in result if len(line) >= 2)
        except Exception as exc:
            print(f"[警告] OCR 识别《{path.name}》第 {page_index + 1} 页时出错：{exc}")
            return ""

    # ------------------------------------------------------------------ #
    # Word 解析
    # ------------------------------------------------------------------ #
    def _load_docx(self, path: Path) -> List[Document]:
        """用 python-docx 提取全部非空段落，合并为一个 Document。"""
        try:
            import docx
        except ImportError as exc:
            raise ImportError("缺少 python-docx 依赖，请先执行：pip install python-docx") from exc

        doc = docx.Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            print(f"[警告] 《{path.name}》中未读取到有效段落。")
            return []

        # Word 文档没有"页"的概念，整体作为一个 Document，
        # 交给下游 TextSplitter 按语义切分即可。
        return [Document(
            page_content="\n".join(paragraphs),
            metadata={"source": path.name, "page": None, "doc_type": "docx"},
        )]

    # ------------------------------------------------------------------ #
    # TXT 解析
    # ------------------------------------------------------------------ #
    def _load_txt(self, path: Path) -> List[Document]:
        """直接读取纯文本文件，自动尝试 utf-8 / gbk 两种常见编码。"""
        text = None
        for encoding in ("utf-8", "gbk"):
            try:
                text = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue
        if text is None:
            raise ValueError(f"无法识别文件编码：{path.name}（仅尝试过 utf-8 / gbk）")

        text = text.strip()
        if not text:
            print(f"[警告] 文本文件《{path.name}》内容为空。")
            return []
        return [Document(
            page_content=text,
            metadata={"source": path.name, "page": None, "doc_type": "txt"},
        )]


if __name__ == "__main__":
    # 简单自测：python document_loader.py <文件或目录路径>
    import sys

    if len(sys.argv) < 2:
        print("用法：python document_loader.py <文件或目录路径>")
        sys.exit(0)

    target = Path(sys.argv[1])
    loader = DocumentLoader()
    results = loader.load_directory(target) if target.is_dir() else loader.load(target)
    print(f"\n共解析出 {len(results)} 个 Document，预览第一条：")
    if results:
        first = results[0]
        print(f"metadata = {first.metadata}")
        print(f"正文前 100 字 = {first.page_content[:100]}")
